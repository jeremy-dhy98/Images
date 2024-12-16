import docx
import win32com.client as win32
from docx.shared import Inches
from docx2pdf import convert
from pathlib import Path

# Define path for saving the Word document
path = Path.home().joinpath("Desktop", "CSV", "logos.docx")  # Use '.docx' extension

# Create a new Word document
doc = docx.Document()
doc.add_heading("LOGO DESIGNING", level=0)

# Add introductory text
doc.add_paragraph(
    """Are you an individual, a medium business enterprise, or an established company looking for a way to distinguish your product's identity among other similar products? 
"JEREMY ARTS AND DESIGN" is here to help you sort out this problem. \n
Below is one of our designs. Reach out and boost your business by  letting us craft professional Logos curated to meet your needs."""
)

# Paths to logo images
paths = [
    r"logo.png", r"jeremy_logo.png", r"enhanced_jeremy_logo.png",
    r"jeremy_logo_black_text.png", r"jeremy_logo_with_circle.png",
    r"jeremy_logo_with_small_stars.png", r"jeremy_logo_with_enhancements.png",
    r"jeremy_logo_with_square.png", r"jeremy_logo_with_enhancements&.png",
    r"jeremy_logo_with_circle_text_2.png", r"jeremy_logo_1280x720.png",
    r"jeremy_logo_updated_2.png"
]

# Add images to the document
for img_path in paths:
    if Path(img_path).is_file():  # Ensure the image file exists
        doc.add_picture(img_path, width=Inches(3), height=Inches(4))
    else:
        print(f"Image not found: {img_path}")

# Add contact information
contact_paragraph = doc.add_paragraph("CONTACT US ON: ")
contact_paragraph.add_run("\nEmail: mulwajeremy24@gmail.com\nPhone: 0713916894").bold = True

# Save the document
doc.save(path)
print(f"Document saved at: {path}")

# Convert to pdf
# Path to the Word file
input_file = str(Path(Path.home() / "Desktop" / "CSV" / "logos.docx"))
output_file = str(Path(Path.home() / "Desktop" / "CSV" / "logos.pdf"))

# Convert Word document to PDF
convert(input_file, output_file)

print(f"PDF created at {output_file}.")

# OTHER OPTIONS--> Make a docx uneditable
# Path to the generated Word file
file_path = str(Path(Path.home() / "Desktop" / "CSV" / "logos.docx"))

# Open Word application
word = win32.Dispatch("Word.Application")
word.Visible = False

# Open the document
doc = word.Documents.Open(file_path)

# Apply restrictions (Read-Only)
doc.Protect(2, Password="JeremyLogo")  # 2 = wdAllowOnlyReading

# Save and close the document
doc.Save()
doc.Close()
word.Quit()

print(f"Document at {file_path} is now read-only.")



